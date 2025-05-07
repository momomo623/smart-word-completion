"""文档填充器."""

import docx
from docx import Document
from loguru import logger

from src.data.models import PlaceholderInfo


class DocumentFiller:
    """文档填充器."""
    
    def __init__(self):
        """初始化文档填充器."""
    
    def fill_neutral_term(
        self, doc: Document, placeholder: PlaceholderInfo, neutral_term: str
    ) -> None:
        """用中性词填充占位符.
        
        Args:
            doc: Document对象
            placeholder: 占位符信息
            neutral_term: 中性词
        """
        try:
            # 根据占位符类型选择不同的填充方法
            if placeholder.placeholder_type == "table":
                self._fill_table_placeholder(doc, placeholder, neutral_term)
            else:
                self._fill_text_placeholder(doc, placeholder, neutral_term)
            
            # 更新占位符的中性词
            placeholder.neutral_term = neutral_term
            
        except Exception as e:
            logger.error(f"填充中性词失败: {e}")
            raise ValueError(f"填充中性词失败: {e}")
    
    def _fill_text_placeholder(
        self, doc: Document, placeholder: PlaceholderInfo, neutral_term: str
    ) -> None:
        """填充文本类型占位符（段落中），支持跨 run 替换，兼容多种检测器."""
        para = doc.paragraphs[placeholder.paragraph_index]
        runs = para.runs

        # 构造替换内容
        if neutral_term == "???":
            replacement = ('{{' + neutral_term + '}}')
            highlighted = True
        else:
            replacement = ('{{' + neutral_term + '}}')
            highlighted = False
        
        # 优先处理colon_field和colon_field_space类型
        if self._fill_colon_field_like(para, runs, placeholder, replacement):
            return
        # 优先跨 run 替换
        if self._try_replace_cross_run(runs, placeholder, replacement, highlighted):
            return

        # 否则，run 内替换
        self._replace_in_single_run(para, runs, placeholder, replacement, highlighted)

    def _try_replace_cross_run(self, runs, placeholder, replacement, highlighted) -> bool:
        """尝试跨 run 替换占位符，成功返回 True，否则 False"""
        # 拼接所有 run 的文本，记录每个 run 的起止位置
        run_ranges = []
        full_text = ''
        for run in runs:
            start = len(full_text)
            full_text += run.text
            end = len(full_text)
            run_ranges.append((start, end))

        target = placeholder.raw_text if placeholder.raw_text else None
        if not target:
            return False
        idx = full_text.find(target)
        if idx == -1:
            return False
        idx_end = idx + len(target)

        # 判断是否只在一个run内
        run_start = run_end = None
        for i, (start, end) in enumerate(run_ranges):
            if idx >= start and idx < end:
                run_start = i
            if idx_end > start and idx_end <= end:
                run_end = i
        if run_start is not None and run_end is not None and run_start == run_end:
            # 只在一个run内，交给单run替换
            return False

        # 否则，执行跨run替换
        replaced = False
        for i, (start, end) in enumerate(run_ranges):
            if end <= idx or start >= idx_end:
                continue
            run = runs[i]
            # 头 run
            if idx >= start and idx < end:
                rel_start = idx - start
                run.text = run.text[:rel_start] + replacement
                if highlighted:
                    run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            # 尾 run
            elif idx_end > start and idx_end <= end:
                rel_end = idx_end - start
                run.text = run.text[rel_end:]
            # 中间 run（完全被替换）
            elif idx < start and idx_end > end:
                run.text = ''
            replaced = True
        if replaced:
            logger.info(f"已将 '{target}' 跨 run 替换为 '{replacement}'")
        return replaced

    def _replace_in_single_run(self, para, runs, placeholder, replacement, highlighted):
        """run 内替换，兼容下划线空格、llm_detected等类型，优先用raw_text精准替换"""
        run_idx = placeholder.run_index if placeholder.run_index < len(runs) else 0
        run = runs[run_idx]

        # 优先用raw_text进行替换
        target = placeholder.raw_text if placeholder.raw_text else None
        
       
            
        if target and target in run.text:
            new_text = run.text.replace(target, replacement, 1)
            run.text = new_text
            if highlighted:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            logger.info(f"已将 '{target}' 替换为 '{replacement}' (run 内)")
            return

        # fallback: 兼容不同类型的占位符
        if placeholder.placeholder_type == "underline":
            import re
            underline_match = re.search(r"_{5,}", run.text)
            if underline_match:
                original_text = underline_match.group(0)
            else:
                original_text = run.text
        elif placeholder.placeholder_type == "underline_space":
            original_text = run.text
        elif placeholder.placeholder_type == "llm_detected":
            original_text = run.text
        else:
            original_text = run.text

        new_text = run.text.replace(original_text, replacement)
        run.text = new_text

        if highlighted:
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

        logger.info(f"已将 '{original_text}' 替换为 '{replacement}' (run 内 fallback)")
    
    def _fill_table_placeholder(
        self, doc: Document, placeholder: PlaceholderInfo, neutral_term: str
    ) -> None:
        """填充表格中的占位符.
        
        对于表格占位符，不使用大模型生成中性词，直接使用表头作为占位符内容。
        为了区分同一列的不同行，在占位符中添加行号作为索引，例如 {{访视时间1}}、{{访视时间2}}。
        
        Args:
            doc: Document对象
            placeholder: 占位符信息
            neutral_term: 中性词（对于表格，这个参数会被忽略）
        """
        # 解析表格索引（使用负值表示）
        table_idx = abs(placeholder.paragraph_index) - 1
        if table_idx >= len(doc.tables):
            logger.error(f"表格索引超出范围: {table_idx} >= {len(doc.tables)}")
            return
        
        table = doc.tables[table_idx]
        
        # 从run_index解析行列信息
        row_idx = placeholder.run_index // 100
        cell_idx = placeholder.run_index % 100
        
        if row_idx >= len(table.rows) or cell_idx >= len(table.rows[0].cells):
            logger.error(f"表格单元格索引超出范围: [{row_idx}, {cell_idx}]")
            return
        
        # 获取单元格
        cell = table.rows[row_idx].cells[cell_idx]
        
        # 对于表格占位符，直接使用占位符的text（即表头）作为内容
        header_text = placeholder.text
        
        # 准备替换文本 - 使用双花括号包裹表头，并添加行号索引
        # 行索引从1开始，更符合用户习惯
        replacement = "{{" + header_text + str(row_idx) + "}}"
        
        # 设置单元格文本
        cell.text = replacement
        
        # 记录使用的表头作为中性词，包含行号
        placeholder.neutral_term = header_text + str(row_idx)
        
        logger.debug(f"已在表格[{table_idx}]的单元格[{row_idx}, {cell_idx}]填入表头占位符: '{replacement}'")

    def _fill_colon_field_like(self, para, runs, placeholder, replacement):
        """
        处理colon_field和colon_field_space类型的回填，返回True表示已处理。
        - colon_field: 在冒号后插入replacement，可能新建run。
        - colon_field_space: 在冒号后第一个非空字符前插入replacement，不新建run。
        Args:
            para: 当前段落对象
            runs: 当前段落的runs列表
            placeholder: 占位符信息对象
            replacement: 要插入的内容（如{{xxx}}）
        Returns:
            bool: True表示已处理，False表示未处理
        """
        if placeholder.placeholder_type not in ("colon_field", "colon_field_space"):
            return False
        run_idx = placeholder.run_index if placeholder.run_index < len(runs) else len(runs) - 1
        # 找到冒号所在run及其在run内的位置
        colon_pos = None
        for i, run in enumerate(runs):
            idx = run.text.rfind('：')
            if idx != -1:
                colon_pos = (i, idx)
                break
        if colon_pos:
            i, idx = colon_pos
            run = runs[i]
            if placeholder.placeholder_type == "colon_field":
                # colon_field: 在冒号后插入replacement，可能新建run
                before = run.text[:idx+1]
                after = run.text[idx+1:]
                run.text = before
                new_run = para.add_run(replacement)
                if after:
                    para.add_run(after)
                return True
            elif placeholder.placeholder_type == "colon_field_space":
                # colon_field_space: 在冒号后第一个非空字符前插入replacement，不新建run
                after_colon = run.text[idx+1:]
                offset = 0
                for offset, ch in enumerate(after_colon):
                    if not ch.isspace():
                        break
                insert_pos = idx + 1 + offset
                run.text = run.text[:insert_pos] + replacement + run.text[insert_pos:]
                return True
        else:
            # 没找到冒号，直接在当前run后插入
            para.add_run(replacement)
            return True
        return False