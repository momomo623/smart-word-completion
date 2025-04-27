"""下划线空格占位符检测器."""

from typing import List, Tuple, Set, Dict

from docx import Document
from docx.enum.text import WD_UNDERLINE
from loguru import logger

from src.data.models import PlaceholderInfo
from src.data.placeholder_detector.base_detector import PlaceholderDetector
from src.data.context_extractor import ContextExtractor


class UnderlineSpaceDetector(PlaceholderDetector):
    """下划线空格占位符检测器.
    
    检测Word文档中带有下划线格式的空格，这通常用作需要填写的字段。
    可以检测两种情况：
    1. 尾部下划线空格 - 文本后带有下划线的空格
    2. 中间下划线空格 - 文本中间带有下划线的空格，如 "文字1{空格下划线}文字2"
    """
    
    def __init__(self, context_window: int = 100):
        """初始化下划线空格占位符检测器.
        
        Args:
            context_window: 上下文窗口大小
        """
        self.context_extractor = ContextExtractor(context_window)
    
    def detect(self, doc: Document) -> List[PlaceholderInfo]:
        """检测文档中的下划线空格占位符.
        
        Args:
            doc: Document对象
            
        Returns:
            占位符信息列表
        """
        placeholders = []
        full_text = "\n".join(para.text for para in doc.paragraphs)
        
        # 使用字典记录每个段落中的占位符位置，以便合并连续的占位符
        # 格式: {段落索引: [(开始位置, 结束位置, 文本块索引), ...]}
        placeholder_positions: Dict[int, List[Tuple[int, int, int]]] = {}
        
        # 遍历所有段落和文本块
        for para_idx, para in enumerate(doc.paragraphs):
            # 检查段落中是否有下划线格式的文本块
            has_underlined_run = False
            for run in para.runs:
                if run.font.underline is not None and run.font.underline == WD_UNDERLINE.SINGLE:
                    has_underlined_run = True
                    break
                    
            if not has_underlined_run:
                continue  # 如果段落中没有下划线格式的文本块，跳过此段落
                
            para_text = para.text
            
            # 初始化段落的占位符位置列表
            if para_idx not in placeholder_positions:
                placeholder_positions[para_idx] = []
                
            for run_idx, run in enumerate(para.runs):
                # 检查是否有下划线
                if run.font.underline is not None and run.font.underline == WD_UNDERLINE.SINGLE:
                    # 分析两种情况：尾部下划线空格和中间下划线空格
                    underline_positions = self._find_underline_space_positions(run.text)
                    
                    if underline_positions:
                        # 获取文本块在段落中的起始位置
                        run_start_pos = self._get_run_start_position(para, run_idx)
                        
                        for start_pos, end_pos in underline_positions:
                            # 获取下划线空格在段落中的绝对位置
                            absolute_start = run_start_pos + start_pos
                            absolute_end = run_start_pos + end_pos
                            
                            # 提取前后文本用于上下文验证
                            placeholder_text = para_text[absolute_start:absolute_end]
                            
                            # 确保占位符文本不为空，且只包含空格
                            if not placeholder_text or not placeholder_text.strip() == '':
                                continue
                                
                            # 记录位置信息
                            placeholder_positions[para_idx].append((absolute_start, absolute_end, run_idx))
        
        # 合并每个段落中连续的占位符位置
        for para_idx, positions in placeholder_positions.items():
            # 按照开始位置排序
            sorted_positions = sorted(positions, key=lambda x: x[0])
            
            # 合并连续的位置
            merged_positions = []
            if sorted_positions:
                current_start, current_end, current_run_idx = sorted_positions[0]
                
                for next_start, next_end, next_run_idx in sorted_positions[1:]:
                    # 如果当前结束位置等于下一个开始位置，合并它们
                    if current_end >= next_start - 1:  # 允许1个字符的小间隔
                        current_end = max(current_end, next_end)
                    else:
                        # 否则，保存当前合并的位置
                        merged_positions.append((current_start, current_end, current_run_idx))
                        current_start, current_end, current_run_idx = next_start, next_end, next_run_idx
                
                # 添加最后一个合并的位置
                merged_positions.append((current_start, current_end, current_run_idx))
            
            # 根据合并后的位置创建占位符
            para = doc.paragraphs[para_idx]
            para_text = para.text
            
            for start_pos, end_pos, run_idx in merged_positions:
                # 提取占位符文本
                placeholder_text = para_text[start_pos:end_pos]
                
                # 从全文提取上下文
                try:
                    full_context_pos = full_text.find(para_text)
                    if full_context_pos != -1:
                        placeholder_full_pos = full_context_pos + start_pos
                        before_text = full_text[max(0, placeholder_full_pos - 100):placeholder_full_pos]
                        after_text = full_text[placeholder_full_pos + len(placeholder_text):
                                            placeholder_full_pos + len(placeholder_text) + 100]
                    else:
                        # 如果在全文中找不到段落，就使用段落内的上下文
                        before_text = para_text[max(0, start_pos - 50):start_pos]
                        after_text = para_text[end_pos:min(len(para_text), end_pos + 50)]
                except Exception as e:
                    logger.warning(f"提取上下文失败: {e}")
                    before_text = ""
                    after_text = ""
                
                # 创建占位符信息
                placeholder = PlaceholderInfo(
                    text="下划线空格占位符",
                    paragraph_index=para_idx,
                    run_index=run_idx,
                    before_text=before_text.strip(),
                    after_text=after_text.strip(),
                    placeholder_type="underline_space",
                    line_text=para_text,
                )
                
                placeholders.append(placeholder)
                # logger.debug(f"找到下划线空格占位符: 段落={para_idx}, 位置={start_pos}-{end_pos}, 长度={len(placeholder_text)}")
        
        logger.info(f"UnderlineSpaceDetector 共检测到 {len(placeholders)} 个占位符")
        return placeholders
    
    def _find_underline_space_positions(self, text: str) -> List[Tuple[int, int]]:
        """查找文本中下划线空格的位置.
        
        Args:
            text: 文本内容
            
        Returns:
            下划线空格的起始和结束位置列表 [(start_pos, end_pos), ...]
        """
        positions = []
        
        # 查找尾部下划线空格
        stripped_text = text.rstrip(' ')
        trailing_spaces = len(text) - len(stripped_text)
        if trailing_spaces > 0:
            positions.append((len(stripped_text), len(text)))
        
        # 查找中间下划线空格
        if ' ' in text[:-trailing_spaces if trailing_spaces > 0 else None]:
            # 查找所有空格位置
            space_positions = [i for i, char in enumerate(text) if char == ' ']
            
            # 排除尾部已处理的空格
            if trailing_spaces > 0:
                space_positions = [pos for pos in space_positions if pos < len(stripped_text)]
            
            # 对连续空格进行分组
            if space_positions:
                start_pos = space_positions[0]
                for i in range(1, len(space_positions)):
                    if space_positions[i] != space_positions[i-1] + 1:
                        # 非连续空格，记录之前的空格组
                        positions.append((start_pos, space_positions[i-1] + 1))
                        start_pos = space_positions[i]
                
                # 处理最后一组空格
                positions.append((start_pos, space_positions[-1] + 1))
        
        return positions
    
    def _get_run_start_position(self, paragraph, run_idx: int) -> int:
        """获取文本块在段落中的起始位置.
        
        Args:
            paragraph: 段落对象
            run_idx: 文本块索引
            
        Returns:
            文本块在段落中的起始位置
        """
        position = 0
        for i in range(run_idx):
            position += len(paragraph.runs[i].text)
        return position