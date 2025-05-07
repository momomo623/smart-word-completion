import sys
from pathlib import Path
from typing import List
from docx import Document
from docx.document import Document as DocxDocument
from docx import Document as DocxDocumentClass
from docx.oxml import OxmlElement
from loguru import logger

sys.path.append(str(Path(__file__).parent.parent))
from src.data.models import PlaceholderInfo
from src.data.document_handler import DocumentHandler

class TableCellProcessor:
    """表格单元格处理器：将每个cell视为迷你文档，检测并返回带表格信息的占位符。"""
    def __init__(self):
        self.handler = DocumentHandler()

    def process_table_cells(self, doc: Document) -> List[PlaceholderInfo]:
        placeholders = []
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # 构造一个仅包含cell.paragraphs的临时Document对象
                    temp_doc = self._create_temp_doc_from_paragraphs(cell.paragraphs)
                    # 输出temp_doc结构
                    for i, para in enumerate(temp_doc.paragraphs):
                        print(f"段落{i}: {para.text.strip()}")
                        for j, run in enumerate(para.runs):
                            print(f"  └─ run{j}: '{run.text}' []")
                    
                    cell_placeholders = self.handler.find_placeholders(temp_doc)
                    for p in cell_placeholders:
                        p.table_index = table_idx
                        p.row_index = row_idx
                        p.col_index = col_idx
                        # cell内段落索引
                        # p.paragraph_index = para_idx  # 保持原有索引
                        placeholders.append(p)
        logger.info(f"表格处理器共检测到 {len(placeholders)} 个占位符")
        return placeholders

    def _create_temp_doc_from_paragraphs(self, paragraphs) -> Document:
        """从一组段落对象构造一个临时Document对象，仅用于检测。"""
        temp_doc = DocxDocumentClass()
        # 清空默认段落
        temp_doc._body.clear_content()
        for para in paragraphs:
            # 复制段落xml
            temp_doc._body._element.append(para._element)
        return temp_doc

# 测试代码
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="表格单元格占位符检测测试")
    parser.add_argument("input", help="输入Word文档路径")
    args = parser.parse_args()
    doc = Document(args.input)
    processor = TableCellProcessor()
    placeholders = processor.process_table_cells(doc)
    for p in placeholders:
        print(p)
