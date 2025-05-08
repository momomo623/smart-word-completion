import os
import re
from typing import List, Dict, Any
from docx import Document
from pocketflow import AsyncFlow, Node, AsyncParallelBatchNode
from utils.llm_client import LLMClient
from utils.document_io import read_docx, save_docx
from loguru import logger
# 注意：本流程包含异步节点，必须用await flow.run_async(shared)调用。

class BatchReadDocNode(Node):
    def prep(self, shared):
        return shared["doc_paths"]
    def exec(self, doc_paths: List[str]):
        docs = [read_docx(path) for path in doc_paths]
        return docs
    def post(self, shared, prep_res, exec_res):
        shared["docs"] = exec_res
        return "default"

class DispatchNode(Node):
    """
    类型分发节点，将所有段落、单列表格单元格、多列表格行分发为任务列表。
    """
    def prep(self, shared):
        return shared["docs"]
    def exec(self, docs):
        tasks_para = []
        tasks_table_row = []
        for doc_idx, doc in enumerate(docs):
            # 段落
            for para_idx, para in enumerate(doc["paragraphs"]):
                if para.text.strip() == "":
                    continue
                tasks_para.append({
                    "doc_idx": doc_idx,
                    "para_idx": para_idx,
                    "para": para,
                    "type": "paragraph"
                })
            # 表格
            for table_idx, table in enumerate(doc["tables"]):
                n_cols = len(table.columns)
                if n_cols == 1:
                    # 单列表格，逐格处理
                    for row_idx, row in enumerate(table.rows):
                        cell = row.cells[0]
                        for cell_para_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip() == "":
                                continue
                            tasks_para.append({
                                "doc_idx": doc_idx,
                                "table_id": table_idx,
                                "row_id": row_idx,
                                "col_id": 0,
                                "para_id": cell_para_idx,
                                "para": para,
                                "type": "table_cell"
                            })
                else:
                    # 多列表格，按行整体处理，跨行同列去重（cell对象唯一）
                    seen_tc_ids = set()
                    for row_idx, row in enumerate(table.rows):
                        row_cells = []
                        for col_idx, cell in enumerate(row.cells):
                            tc_id = id(cell._tc)
                            if tc_id in seen_tc_ids:
                                continue  # 跳过合并行导致的重复cell
                            seen_tc_ids.add(tc_id)
                            cell_text = "\n".join([p.text for p in cell.paragraphs if p.text.strip()])
                            row_cells.append({
                                "col_id": col_idx,
                                "cell_text": cell_text,
                                "cell": cell,
                                "real_col_id": col_idx
                            })
                        tasks_table_row.append({
                            "doc_idx": doc_idx,
                            "table_id": table_idx,
                            "row_id": row_idx,
                            "row_cells": row_cells,
                            "type": "table_row"
                        })
        return {"para_tasks": tasks_para, "table_row_tasks": tasks_table_row}
    def post(self, shared, prep_res, exec_res):
        shared["para_tasks"] = exec_res["para_tasks"]
        shared["table_row_tasks"] = exec_res["table_row_tasks"]
        # 按任务类型分发action
        actions = []
        # if exec_res["para_tasks"]:
        #     actions.append("para")
        # if exec_res["table_row_tasks"]:
        #     actions.append("table_row")
        # 如果都有，优先para，PocketFlow会依次分支
        # return actions[0] if actions else None
        
        # 格式化输出para_tasks和table_row_tasks
        # for task in exec_res["para_tasks"]:
        #     logger.info(f"\n段落任务：\n{task}\n")
        for task in exec_res["table_row_tasks"]:
            print("-"*80)
            print(f"行{task['row_id']}")
            for col in  task["row_cells"]:
                print(f"列{col['col_id']}   {col['cell_text']}")
        return "default"

class ParaLLMFillNode(AsyncParallelBatchNode):
    def __init__(self):
        super().__init__()
        self.llm_client = LLMClient()
    @staticmethod
    def extract_neutral_term(text):
        match = re.search(r"\{\{([^{}]+)\}\}", text)
        if match:
            return match.group(1)
        match = re.search(r"\{([^{}]+)\}", text)
        if match:
            return match.group(1)
        return ""
    async def prep_async(self, shared):
        return shared.get("para_tasks", [])
    async def exec_async(self, item):
        para = item["para"]
        paragraph_text = para.text
        runs_text = "\n".join([f'  run{j}: "{run.text.strip()}"' for j, run in enumerate(para.runs)])
        logger.info(f"\n段落：{paragraph_text}\n每个run：{runs_text}")
        prompt = f"""
你是一名专业的内容分析助手。你的任务是审查下方Word文档的段落结构，判断每个run是否包含占位符或需要人工填写的内容（如姓名、日期、专业等字段）。

**任务要求**
- 如果所有run都不需要填写，将need_fill设为false，fill_list设为空字典。
- 如果有run需要填写，将need_fill设为true，并在fill_list中列出所有run序号及对应run的回填后内容（即将占位符替换为中性词后的完整文本，格式如：{{中性词}}）。
- 前序字段继承原则：优先继承前文的字段名称（如"姓名:____"直接继承"姓名"）。
- 后序字段继承原则：如果前序字段不明确，可结合后文字段（如"xxxx专业申请"应输出"专业名称"）。
- 中性描述词总结：如果没有前序和后续字段名称，根据上下文提取最贴切的名词短语。
- 尤其关注冒号，如果有冒号大概率需要填写中性词。
- 不要修改原始文本，只能替换占位符或者添加中性词
- 安全边界：如上下文字段不明确，请返回"???"。
- 在分析过程中，请逐步思考，但每个步骤的描述尽量简洁（不超过10个字）。
- 使用分隔符"####"来区分思考过程与最终答案，"####"后直接输出YAML。

**输出格式**
思考1:（不超过10个字）
思考2:（不超过10个字）
####
```yaml
need_fill: false/true
fill_list:
  run_index: run_filled_text
```

**输入示例**
完整段落：
由 (xxxxxx公司) 申在我院xxx专业申请开展的一项名为"（xxxxxx临床试验） "的临床研究。
每个run：
  run0: "由 ("
  run1: "xxxxxx"
  run2: "公司) 申在我院xxx专业申请开展的一项名为"（"
  run3: "xxxxxx"
  run4: "临床试验） "的临床研究。"

**输出示例**
思考1:（不超过10个字）
思考2:（不超过10个字）
####
```yaml
need_fill: true
fill_list:
  1: "{{公司名称}}"
  2: "公司) 申在我院{{专业名称}}专业申请开展的一项名为"（"
  3: "{{临床试验名称}}"
  
**错误示例**
原始run0: '联系电话/传真/手机：'
错误输出: "{{联系电话}}/{{传真}}/{{手机}}："
正确输出: "联系电话/传真/手机：{{联系电话/传真/手机}}"
```

## 现在，请分析下方段落：
完整段落：
{paragraph_text}
每个run：
{runs_text}
"""
        response = self.llm_client.chat_completion(user_message=prompt)
        yaml_data = self.llm_client.parse_yaml(response)
        logs = []
        if yaml_data.get("need_fill"):
            fill_list = yaml_data.get("fill_list") or {}
            
            for run_idx, run_filled_text in fill_list.items():
                try:
                    idx = int(run_idx)
                    para.runs[idx].text = run_filled_text
                    neutral = self.extract_neutral_term(run_filled_text)
                    log = {
                        "type": item["type"],
                        "doc_idx": item["doc_idx"],
                        "para_idx": item.get("para_idx"),
                        "table_id": item.get("table_id"),
                        "row_id": item.get("row_id"),
                        "col_id": item.get("col_id"),
                        "cell_para_id": item.get("para_id"),
                        "run_id": idx,
                        "neutral_term": neutral
                    }
                    logs.append(log)
                except Exception as e:
                    logs.append({"error": str(e), **item})
        return {
            "type": item["type"],
            "doc_idx": item["doc_idx"],
            "para_idx": item.get("para_idx"),
            "table_id": item.get("table_id"),
            "row_id": item.get("row_id"),
            "col_id": item.get("col_id"),
            "cell_para_id": item.get("para_id"),
            "filled_text": para.text,
            "logs": logs
        }
    async def post_async(self, shared, prep_res, exec_res_list):
        shared["para_results"] = exec_res_list
        return "default"

class TableRowLLMFillNode(AsyncParallelBatchNode):
    def __init__(self):
        super().__init__()
        self.llm_client = LLMClient()
    async def prep_async(self, shared):
        # 将docs对象一并传递到每个任务item中
        docs = shared["docs"]
        tasks = shared.get("table_row_tasks", [])
        for task in tasks:
            task["docs"] = docs
        return tasks
    async def exec_async(self, item):
        row_cells = item["row_cells"]
        virtual2real = {}
        row_texts = []
        for v_idx, cell in enumerate(row_cells):
            virtual2real[str(v_idx)] = cell["real_col_id"]
            cell_obj = cell["cell"]
            run_lines = []
            for para_id, para in enumerate(cell_obj.paragraphs):
                for run_id, run in enumerate(para.runs):
                    if run.text.strip():
                        run_lines.append(f"  - para_id={para_id}, run_id={run_id}, text=\"{run.text.strip()}\"")
            if run_lines:
                cell_block = f"列{v_idx}:\n" + "\n".join(run_lines)
            else:
                cell_block = f"列{v_idx}: <空>"
                
            row_texts.append(cell_block)
        logger.info(f"\n结构化表格行输入：\n{chr(10).join(row_texts)}\n")
        prompt = f"""
你是一名专业的内容分析助手。你将收到一个Word表格的一行内容，每一列都详细列出了所有run的信息（para_id, run_id, text）。
你的任务是分析下方Word表格的一行内容，判断每一列的每个run是否包含占位符或需要人工填写的内容（如姓名、日期、专业等字段），并输出结构化结果。

**任务要求**
- 你需要结合本行所有列的内容，判断每个run是否需要填写内容。将需要填写的内容总结为一个中性词，和原始文本作为run_filled_text一起输出，(如"原始文本:{{中性词}}")。
- 如果某个run不需要修改，则无需输出
- 如果某格内容为空，有可能依赖于上一格，请根据上一格的文本判断是否需要填写
- 如果2列依赖1列，则2列需要填写，1列不需要填写任何内容
- 尤其关注冒号，如果有冒号大概率需要填写中性词
- 如果某格内容为选项框，不要填写
- 不要修改原始文本，只能替换占位符或者添加中性词
- 在分析过程中，请逐步思考，但每个步骤的描述尽量简洁（不超过10个字）。
- 使用分隔符"####"来区分思考过程与最终答案，"####"后直接输出YAML。

**输出格式**
思考1:（不超过10个字）
思考2:（不超过10个字）
####
```yaml
- col_id: 列序号
  para_id: 段落序号
  run_id: 段落run序号
  run_filled_text: 回填后的文本
```

**输入示例**
列0:
  - para_id=0, run_id=0, text="方案信息"
列1:
  - para_id=0, run_id=0, text="方案版本号及日期"
列2:
  <空>
列3:
  - para_id=0, run_id=0, text="知情同意书版本号日期"
列4:
  - para_id=0, run_id=0, text="/"
列5:
  - para_id=0, run_id=0, text="项目名称："

**输出示例**
思考1:列0没有明显的占位符，也不依赖于其他列，不需要填写中性词
思考2:列1没有明显占位符，列2依赖列1
思考3:列4依赖于列3
思考4:列5有明显的占位符，且没有其他列依赖列5，需要填写中性词
####
```yaml
- col_id: 2
  para_id: 0
  run_id: 0
  run_filled_text: "{{方案版本号及日期}}"
- col_id: 4
  para_id: 0
  run_id: 0
  run_filled_text: "{{知情同意书版本号日期}}"
- col_id: 5
  para_id: 0
  run_id: 0
  run_filled_text: "项目名称：{{项目名称}}"
```


## 现在，请分析下方表格行：
{chr(10).join(row_texts)}
"""
        response = self.llm_client.chat_completion(user_message=prompt)
        yaml_data = self.llm_client.parse_yaml(response)
        logs = []
        docs = item["docs"]  # 通过prep_async传递进来
        doc_idx = item["doc_idx"]
        for entry in yaml_data:
            try:
                v_col_id = entry["col_id"]  # 大模型返回的虚拟编号
                para_id = entry["para_id"]
                run_id = entry["run_id"]
                filled_text = entry["run_filled_text"]
                if not isinstance(filled_text, str):
                    filled_text = str(filled_text)
                real_col_id = virtual2real[str(v_col_id)]
                cell = docs[doc_idx]["tables"][item["table_id"]].rows[item["row_id"]].cells[real_col_id]
                # 1. 确保有段落
                if not cell.paragraphs:
                    para = cell.add_paragraph()
                elif para_id >= len(cell.paragraphs):
                    para = cell.add_paragraph()
                else:
                    para = cell.paragraphs[para_id]
                # 2. 确保有run
                if not para.runs:
                    run = para.add_run()
                elif run_id >= len(para.runs):
                    run = para.add_run()
                else:
                    run = para.runs[run_id]
                run.text = filled_text
                log = {
                    "type": "table_row",
                    "doc_idx": doc_idx,
                    "table_id": item["table_id"],
                    "row_id": item["row_id"],
                    "col_id": real_col_id,
                    "para_id": para_id,
                    "run_id": run_id,
                    "run_filled_text": filled_text
                }
                logs.append(log)
            except Exception as e:
                logs.append({"error": str(e), **entry})
        return {
            "type": "table_row",
            "doc_idx": item["doc_idx"],
            "table_id": item["table_id"],
            "row_id": item["row_id"],
            "logs": logs
        }
    async def post_async(self, shared, prep_res, exec_res_list):
        shared["table_row_results"] = exec_res_list
        return "default"

class MergeResultNode(Node):
    def prep(self, shared):
        return {
            "docs": shared["docs"],
            "para_results": shared.get("para_results", []),
            "table_row_results": shared.get("table_row_results", [])
        }
    def exec(self, data):
        # 此处可实现回填合并、日志合并等
        parse_logs = []
        for res in data["para_results"]:
            if res.get("logs"):
                parse_logs.extend(res["logs"])
        for res in data["table_row_results"]:
            if res.get("logs"):
                parse_logs.extend(res["logs"])
        return {"parse_logs": parse_logs}
    def post(self, shared, prep_res, exec_res):
        shared["parse_logs"] = exec_res["parse_logs"]
        return "default"

class OutputDocNode(Node):
    def prep(self, shared):
        return {
            "docs": shared["docs"],
            "output_paths": shared["output_paths"] if "output_paths" in shared else []
        }
    def exec(self, data):
        docs = data["docs"]
        output_paths = data["output_paths"]
        result_paths = []
        for i, doc_info in enumerate(docs):
            doc = doc_info["doc_obj"]
            out_path = output_paths[i] if i < len(output_paths) else doc_info["doc_path"].replace(".docx", "_filled.docx")
            save_docx(doc, out_path)
            result_paths.append(out_path)
        return result_paths
    def post(self, shared, prep_res, exec_res):
        shared["output_paths"] = exec_res
        return None

# PocketFlow流程组装
def create_llm_detector_flow():
    read_node = BatchReadDocNode()
    dispatch_node = DispatchNode()
    para_node = ParaLLMFillNode()
    table_row_node = TableRowLLMFillNode()
    merge_node = MergeResultNode()
    output_node = OutputDocNode()
    # 串联处理：para_node处理完后再处理table_row_node
    read_node >> dispatch_node >> para_node >> table_row_node >> merge_node >> output_node
    return AsyncFlow(start=read_node)
