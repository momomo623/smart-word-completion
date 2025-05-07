from typing import Dict, Any
from docx import Document

def read_docx(doc_path: str) -> Dict[str, Any]:
    """读取Word文档，返回段落、表格和doc对象。"""
    doc = Document(doc_path)
    paragraphs = [para for para in doc.paragraphs]
    tables = [table for table in doc.tables]
    return {"doc_path": doc_path, "paragraphs": paragraphs, "tables": tables, "doc_obj": doc}


def save_docx(doc: Document, output_path: str):
    """保存Word文档到指定路径。"""
    doc.save(output_path)
